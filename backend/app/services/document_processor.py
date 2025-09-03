from PyPDF2 import PdfReader
from docx import Document
import os
import tiktoken

def extract_text(file_path):
    if file_path.lower().endswith('.pdf'):
        return extract_text_from_pdf(file_path)
    elif file_path.lower().endswith('.docx'):
        return extract_text_from_docx(file_path)
    elif file_path.lower().endswith('.txt'):
        return extract_text_from_txt(file_path)
    else:
        raise ValueError("Unsupported file type for text extraction")

def extract_text_from_pdf(file_path):
    # Basic validation
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    if not file_path.lower().endswith('.pdf'):
        raise ValueError("File must have .pdf extension")
    
    text = ""
    try:
        with open(file_path, "rb") as file:
            reader = PdfReader(file)
            
            # Handle encrypted PDFs
            if reader.is_encrypted:
                reader.decrypt("")  # Try empty password
            
            # Extract text
            for page in reader.pages:
                text += page.extract_text() + "\n"
                
    except Exception as e:
        raise RuntimeError(f"Error processing PDF: {str(e)}")
    
    return text.strip()

def extract_text_from_docx(file_path):
    # Basic validation
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    if not file_path.lower().endswith('.docx'):
        raise ValueError("File must have .docx extension")
    
    text = ""
    try:
        doc = Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
    except Exception as e:
        raise RuntimeError(f"Error processing DOCX: {str(e)}")
    
    return text.strip()

def extract_text_from_txt(file_path):
    # Basic validation
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    if not file_path.lower().endswith('.txt'):
        raise ValueError("File must have .txt extension")
    
    text = ""
    try:
        with open(file_path, "r", encoding="utf-8") as file:
            text = file.read()
    except Exception as e:
        raise RuntimeError(f"Error processing TXT: {str(e)}")
    
    return text.strip()



def preprocess_text(text):
    """
    Clean text while preserving paragraph structure.
    """
    import re
    
    # Step 1: Replace multiple empty lines with double newlines
    cleaned = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
    
    # Step 2: Clean up individual lines and remove single empty lines
    lines = []
    for line in cleaned.splitlines():
        stripped = line.strip()
        if stripped:  # Keep non-empty lines
            lines.append(stripped)
        elif lines and lines[-1]:  # Keep empty line only if it follows content
            lines.append('')
    
    return '\n'.join(lines)

## Text chunking function using tiktoken

def chunk_text(text, max_tokens=300, overlap=30):
    encoding= tiktoken.encoding_for_model("gpt-3.5-turbo")
    tokens=encoding.encode(text)
    chunk_texts=[]
    start=0
    while start < len(tokens):
        end=min(start + max_tokens, len(tokens))
        chunk=encoding.decode(tokens[start:end])
        chunk_texts.append(chunk)
        start += max_tokens - overlap
    return chunk_texts

## metadata extraction functions:

def extract_metadata(file_path):
    if file_path.lower().endswith('.pdf'):
        return extract_metadata_from_pdf(file_path)
    elif file_path.lower().endswith('.docx'):
        return extract_metadata_from_docx(file_path)
    elif file_path.lower().endswith('.txt'):
        return extract_metadata_from_txt(file_path)
    else:
        raise ValueError("Unsupported file type for metadata extraction")

def extract_metadata_from_pdf(file_path):
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    if not file_path.lower().endswith('.pdf'):
        raise ValueError("File must have .pdf extension")
    
    metadata = {}
    try:
        with open(file_path, "rb") as file:
            reader = PdfReader(file)
            doc_info = reader.metadata
            if doc_info:
                for key, value in doc_info.items():
                    clean_key = key.replace('/', '')  # Remove leading slash
                    metadata[clean_key] = value
            metadata['num_pages'] = len(reader.pages)
            metadata['file_size'] = os.path.getsize(file_path)
           
    except Exception as e:
        raise RuntimeError(f"Error extracting PDF metadata: {str(e)}")
    
    return metadata

def extract_metadata_from_docx(file_path):
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    if not file_path.lower().endswith('.docx'):
        raise ValueError("File must have .docx extension")
    
    metadata = {}
    try:
        doc = Document(file_path)
        core_props = doc.core_properties
        metadata = {
            'author': core_props.author,
            'category': core_props.category,
            'comments': core_props.comments,
            'content_status': core_props.content_status,
            'created': core_props.created.isoformat() if core_props.created else None,
            'identifier': core_props.identifier,
            'keywords': core_props.keywords,
            'last_modified_by': core_props.last_modified_by,
            'last_printed': core_props.last_printed.isoformat() if core_props.last_printed else None,
            'modified': core_props.modified.isoformat() if core_props.modified else None,
            'revision': core_props.revision,
            'subject': core_props.subject,
            'title': core_props.title,
            'version': core_props.version,
            'file_size': os.path.getsize(file_path),
        }
    except Exception as e:
        raise RuntimeError(f"Error extracting DOCX metadata: {str(e)}")
    
    return metadata

def extract_metadata_from_txt(file_path):
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    if not file_path.lower().endswith('.txt'):
        raise ValueError("File must have .txt extension")
    
    metadata = {}
    try:
        metadata['file_size'] = os.path.getsize(file_path)
        metadata['created'] = os.path.getctime(file_path)
        metadata['modified'] = os.path.getmtime(file_path)
    except Exception as e:
        raise RuntimeError(f"Error extracting TXT metadata: {str(e)}")
    
    return metadata

